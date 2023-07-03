# ''''''"""
# @author: Nan Jia
# @time: 2023/6/21 14:23
# @desc:
# ''''''"""
#
# from docx import Document
#
# # 创建一个新的Word文档
# document = Document()
#
# # 添加标题
# document.add_heading('撤离交通情况分析报告', level=1)
#
# # 添加报告概述
# document.add_paragraph('报告概述：')
# document.add_paragraph('本报告旨在对撤离交通情况进行深入分析，以揭示交通系统的运行状况和潜在问题。通过综合考虑交通流量、路段拥堵、车辆密度等指标，我们将为改进交通规划和管理提供有价值的见解和建议。')
#
# # 添加正文内容
# document.add_heading('1. 引言', level=2)
# document.add_paragraph('介绍报告目的和背景，突出撤离交通情况的重要性和影响。')
# document.add_paragraph('阐述研究方法和数据来源的选择依据。')
#
# # 添加更多章节和内容...
#
# # 保存Word文档
# document.save('report.docx')


# from docx import Document
# from docx.shared import Pt
# from docx.oxml.ns import nsdecls
# from docx.oxml import parse_xml
#
# # 创建一个新的Word文档
# document = Document()
#
# # 添加段落
# paragraph = document.add_paragraph()
#
# # 设置段落文本
# run = paragraph.add_run('这是一段中文学术的标准字体')
#
# # 设置字体名称和大小
# font = run.font
# font.name = 'SimSun'  # 设置为中文学术的标准字体，如宋体（SimSun）
# font.size = Pt(12)  # 设置字体大小为12磅
#
# # 设置段落的其他样式，如对齐方式、缩进等
#
# # 保存Word文档
# document.save('report1.docx')


from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# 创建一个新的Word文档
document = Document()

# 添加段落
paragraph = document.add_paragraph()

# 设置段落文本
run = paragraph.add_run('这是一段中文学术的标准字体')

# 清除已有的字体样式
run.font.clear()

# 设置字体名称和大小
run.font.name = 'SimSun'  # 设置为中文学术的标准字体，如宋体（SimSun）
run.font.size = Pt(12)  # 设置字体大小为12磅

# 设置段落的其他样式，如对齐方式、缩进等

# 保存Word文档
document.save('report2.docx')
